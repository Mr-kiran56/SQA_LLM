from langchain_text_splitters import TokenTextSplitter
from langchain_community.docstore.document import Document
import os

# For PDF
from langchain_community.document_loaders import PyPDFLoader
# For Word
from docx import Document as DocxDocument

def file_process(file_path):
    """Process any file type (PDF, DOCX, TXT) and split into chunks"""
    print(f"\nProcessing file: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    text_overall = ""

    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
        loader_text = loader.load()
        print(f"Loaded {len(loader_text)} pages from PDF")
        for doc in loader_text:
            text_overall += doc.page_content

    elif ext in [".docx", ".doc"]:
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text_overall += para.text + "\n"
        print(f"Loaded DOCX/DOC with {len(doc.paragraphs)} paragraphs")

    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8") as f:
            text_overall = f.read()
        print(f"Loaded TXT/MD with {len(text_overall.splitlines())} lines")

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Split large text into big chunks
    textdocs = TokenTextSplitter(chunk_size=10000, chunk_overlap=200)
    texts = textdocs.split_text(text_overall)
    print(f"Created {len(texts)} large chunks")

    # Convert to Document objects
    document = [Document(page_content=t) for t in texts]

    # Split into smaller chunks for final processing
    documents_texts = TokenTextSplitter(chunk_size=490, chunk_overlap=400)
    document_chunk = documents_texts.split_documents(document)
    print(f"Created {len(document_chunk)} final chunks")

    return document_chunk
