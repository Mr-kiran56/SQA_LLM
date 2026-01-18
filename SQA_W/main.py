from fastapi import FastAPI, Form, Request, Response, File, Depends, HTTPException, status
from fastapi.responses import RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.encoders import jsonable_encoder
import uvicorn
from fastapi import UploadFile
import os
import aiofiles

from docx import Document
from fastapi.responses import FileResponse
from fastapi import UploadFile, File, Request
# from src.helper import llm_pipeline
from src.file_processing import file_process
from src.llm import generate_questions,save_output
from src.prompt_template import prompt_temp

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")

templates = Jinja2Templates(directory="templates")


@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


GENERATED_TEXT_FILE = "generated_questions.txt"
GENERATED_WORD_FILE = "generated_questions.docx"




import mimetypes

@app.post("/generate")
async def generate_questions_api(
    nums: int = Form(...),
    uploaded_file: UploadFile = File(...)
):
    base_folder = "static"
    os.makedirs(base_folder, exist_ok=True)

    file_path = os.path.join(base_folder, uploaded_file.filename)

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(await uploaded_file.read())

    
    mime_type = uploaded_file.content_type
    print("Uploaded file type:", mime_type)

    chunks = file_process(file_path)

    response = generate_questions(
    document_chunks=chunks,
    filename=uploaded_file.filename,
    num_questions=nums,
    words_per_answer=35
     )


    with open(GENERATED_TEXT_FILE, "w", encoding="utf-8") as f:
        f.write(response)

    return {"status": "generated", "output": response}



@app.get("/download-word")
async def download_word():
    if not os.path.exists(GENERATED_TEXT_FILE):
        raise HTTPException(status_code=404, detail="No generated content")

    doc = Document()
    doc.add_heading("Generated Questions & Answers", level=1)

    with open(GENERATED_TEXT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(GENERATED_WORD_FILE)

    return FileResponse(
        GENERATED_WORD_FILE,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Generated_Questions.docx"
    )





